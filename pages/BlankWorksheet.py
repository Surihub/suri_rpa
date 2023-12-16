import streamlit as st
from kiwipiepy import Kiwi
import docx
from io import BytesIO
import numpy as np

# from transformers import BertTokenizer, BertModel
# import torch

# 페이지 제목 설정
st.title("📝Create Blank Worksheet!")

# 사용법 안내
st.info('###### 사용법\n빈칸을 만들어서 공부하려고 할 때, 원하는 텍스트를 입력 후 특정 단어를 선택해서 빈칸을 만들 수 있습니다.')

# 빈칸을 만들 내용을 입력받음
sample_text = """제1조 ①대한민국은 민주공화국이다.
②대한민국의 주권은 국민에게 있고, 모든 권력은 국민으로부터 나온다.
 제2조 ①대한민국의 국민이 되는 요건은 법률로 정한다.
②국가는 법률이 정하는 바에 의하여 재외국민을 보호할 의무를 진다.
 제3조 대한민국의 영토는 한반도와 그 부속도서로 한다.
 제4조 대한민국은 통일을 지향하며, 자유민주적 기본질서에 입각한 평화적 통일정책을 수립하고 이를 추진한다.
 제5조 ①대한민국은 국제평화의 유지에 노력하고 침략적 전쟁을 부인한다.
②국군은 국가의 안전보장과 국토방위의 신성한 의무를 수행함을 사명으로 하며, 그 정치적 중립성은 준수된다.
 제6조 ①헌법에 의하여 체결ㆍ공포된 조약과 일반적으로 승인된 국제법규는 국내법과 같은 효력을 가진다.
②외국인은 국제법과 조약이 정하는 바에 의하여 그 지위가 보장된다.
 제7조 ①공무원은 국민전체에 대한 봉사자이며, 국민에 대하여 책임을 진다.
②공무원의 신분과 정치적 중립성은 법률이 정하는 바에 의하여 보장된다.
 제8조 ①정당의 설립은 자유이며, 복수정당제는 보장된다.
②정당은 그 목적ㆍ조직과 활동이 민주적이어야 하며, 국민의 정치적 의사형성에 참여하는데 필요한 조직을 가져야 한다.
③정당은 법률이 정하는 바에 의하여 국가의 보호를 받으며, 국가는 법률이 정하는 바에 의하여 정당운영에 필요한 자금을 보조할 수 있다.
④정당의 목적이나 활동이 민주적 기본질서에 위배될 때에는 정부는 헌법재판소에 그 해산을 제소할 수 있고, 정당은 헌법재판소의 심판에 의하여 해산된다.
 제9조 국가는 전통문화의 계승ㆍ발전과 민족문화의 창달에 노력하여야 한다."""

contents = st.text_area("빈칸을 만들 내용을 입력해주세요.", value=sample_text)

@st.cache_data
def show_download():
        # '생성하기' 버튼
    if st.button('빈칸 뚫기 미리보기'):
        # 체크된 명사를 빈칸으로 치환
        for noun in selected_nouns:
            contents = contents.replace(noun, '___'*len(noun))

        # 결과 표시
        st.write(contents)

        # 워드 문서 생성
        doc = docx.Document()
        doc.add_paragraph(contents)
        
        # 워드 파일 다운로드
        with BytesIO() as f:
            doc.save(f)
            f.seek(0)
            st.download_button('워드 파일로 다운로드하기', f, file_name='학습지.docx')


tab1, tab2, tab3 = st.tabs(['문장별로 키워드 선택하기', '수작업[가나다순]', '추천 키워드'])

import pandas as pd
with tab1:

    kiwi = Kiwi()
    # 텍스트를 문장 단위로 분리
    sentences = kiwi.split_into_sents(contents)
    sentences = pd.DataFrame(sentences).text.tolist()

    # 선택된 명사를 저장하는 세션 상태 초기화
    if 'selected_nouns_dict' not in st.session_state:
        st.session_state.selected_nouns_dict = {}

    new_sentences = []

    for i, sentence in enumerate(sentences):
        # 형태소 분석으로 명사 추출
        tokens = kiwi.analyze(sentence)[0][0]
        nouns = [token.form for token in tokens if token.tag == "NNG"]

        # 고유한 키를 가진 멀티셀렉트 생성
        selected_nouns_key = f"selected_nouns_{i}"
        selected_nouns = st.multiselect(f'"{sentence}"', nouns, key=selected_nouns_key)

        # 선택된 명사를 세션 상태에 저장
        st.session_state.selected_nouns_dict[selected_nouns_key] = selected_nouns

        # 현재 문장에서 선택된 명사를 빈칸으로 치환
        for noun in selected_nouns:
            sentence = sentence.replace(noun, '___'*len(noun))

        new_sentences.append(sentence)

    # '생성하기' 버튼
    if st.button('빈칸 뚫기 미리보기', key = 'tab1'):
        # 업데이트된 텍스트 표시
        updated_contents = '\n'.join(new_sentences)
        st.markdown(updated_contents)
        # 워드 문서 생성
        doc = docx.Document()
        doc.add_paragraph(updated_contents)
        
        # 워드 파일 다운로드
        with BytesIO() as f:
            doc.save(f)
            f.seek(0)
            st.download_button('워드 파일로 다운로드', f, file_name='학습지.docx')

with tab2:
    # 형태소 분석
    kiwi = Kiwi()
    tokens = kiwi.analyze(contents)[0][0]

    # 명사를 저장할 집합
    nouns = set()

    # 각 토큰에 대해 명사 추출
    for token in tokens:
        if token.tag == "NNG":
            nouns.add(token.form)
    nouns = sorted(nouns)

    # 명사 선택 위젯
    selected_nouns = st.multiselect('빈칸으로 만들 명사를 선택하세요', list(nouns))
    # '생성하기' 버튼
    if st.button('빈칸 뚫기 미리보기', key = 'tab2'):
        # 체크된 명사를 빈칸으로 치환
        for noun in selected_nouns:
            contents = contents.replace(noun, '___'*len(noun))

        # 결과 표시
        st.write(contents)

        # 워드 문서 생성
        doc = docx.Document()
        doc.add_paragraph(contents)
        
        # 워드 파일 다운로드
        with BytesIO() as f:
            doc.save(f)
            f.seek(0)
            st.download_button('워드 파일로 다운로드', f, file_name='학습지.docx')

with tab3:
    st.write("준비 중 입니다....")
    # # BERT 모델과 토크나이저 초기화
    # tokenizer = BertTokenizer.from_pretrained('bert-base-uncased')
    # model = BertModel.from_pretrained('bert-base-uncased', output_attentions=True)

    # # 사용자 입력 텍스트
    # text = contents

    # # BERT를 사용하여 텍스트 처리
    # inputs = tokenizer(text, return_tensors="pt", add_special_tokens=True)
    # outputs = model(**inputs)
    # attentions = outputs.attentions

    # # Attention 가중치 계산 (마지막 레이어 사용)
    # # 차원 축소를 위해 mean() 대신 squeeze() 사용
    # # Attention 가중치 계산 (마지막 레이어 사용)
    # # 여러 차원을 가진 텐서를 처리하기 위해 mean() 메소드 사용
    # last_layer_attentions = attentions[-1].squeeze(0)
    # word_attentions = last_layer_attentions.mean(dim=0).squeeze()

    # # 가중치가 스칼라가 아닐 경우를 대비해 조건부 처리
    # weighted_tokens = []
    # for token, weight in zip(tokens, word_attentions):
    #     if weight.numel() == 1:  # numel() 메소드는 텐서 내 요소의 총 개수를 반환
    #         weighted_tokens.append((token, weight.item()))
    #     else:
    #         # 텐서에 여러 요소가 있는 경우 평균값 사용
    #         weighted_tokens.append((token, weight.mean().item()))


    # # 상위 10개 중요 단어 추출
    # top_tokens = sorted(weighted_tokens, key=lambda x: x[1], reverse=True)[:10]

    # # 중요 단어 선택 위젯
    # selected_tokens = st.multiselect('중요 단어 선택', [token for token, weight in top_tokens])

    # # '생성하기' 버튼
    # if st.button('빈칸 생성하기'):
    #     # 선택된 토큰을 빈칸으로 치환
    #     for token in selected_tokens:
    #         text = text.replace(token, "____")
    #     st.write(text)

st.success("정답 파일 생성 및 키워드 추천 기능은 준비 중 입니다. ")